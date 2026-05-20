import os
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy

BASE = os.path.dirname(os.path.abspath(__file__))
FILES = [
    "coursework_intro.docx",
    "coursework_ch1.docx",
    "coursework_ch2.docx",
    "coursework_ch3.docx",
    "coursework_conclusion.docx",
    "coursework_refs.docx",
    "coursework_appendix.docx",
]
OUT = os.path.join(BASE, "coursework_full.docx")

def copy_paragraph(src_para, dst_doc):
    """Copy a paragraph with all formatting from src to dst."""
    new_para = dst_doc.add_paragraph()
    
    # Copy paragraph format
    pf_src = src_para.paragraph_format
    pf_dst = new_para.paragraph_format
    
    # Copy alignment
    if pf_src.alignment is not None:
        pf_dst.alignment = pf_src.alignment
    
    # Copy indentation
    if pf_src.left_indent is not None:
        pf_dst.left_indent = pf_src.left_indent
    if pf_src.right_indent is not None:
        pf_dst.right_indent = pf_src.right_indent
    if pf_src.first_line_indent is not None:
        pf_dst.first_line_indent = pf_src.first_line_indent
    
    # Copy spacing
    if pf_src.space_before is not None:
        pf_dst.space_before = pf_src.space_before
    if pf_src.space_after is not None:
        pf_dst.space_after = pf_src.space_after
    if pf_src.line_spacing is not None:
        pf_dst.line_spacing = pf_src.line_spacing
    
    # Copy runs
    for run in src_para.runs:
        new_run = new_para.add_run(run.text)
        r_src = run.font
        r_dst = new_run.font
        
        # Font properties
        if r_src.name is not None:
            r_dst.name = r_src.name
        if r_src.size is not None:
            r_dst.size = r_src.size
        if r_src.bold is not None:
            r_dst.bold = r_src.bold
        if r_src.italic is not None:
            r_dst.italic = r_src.italic
        if r_src.underline is not None:
            r_dst.underline = r_src.underline
        if r_src.color and r_src.color.rgb:
            r_dst.color.rgb = r_src.color.rgb
    
    return new_para

def add_page_number(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run()
        fldChar1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run._r.append(fldChar1)
        
        run2 = para.add_run()
        instrText = run2._r.makeelement(qn('w:instrText'), {})
        instrText.text = ' PAGE '
        run2._r.append(instrText)
        
        run3 = para.add_run()
        fldChar2 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run3._r.append(fldChar2)

# Check all files exist
missing = [f for f in FILES if not os.path.exists(os.path.join(BASE, f))]
if missing:
    print(f"ERROR: Missing files: {missing}")
    exit(1)

# Build the merged document
print("Building merged document...")

# Start with first file as base
merged = Document(os.path.join(BASE, FILES[0]))
print(f"  Added: {FILES[0]}")

# Append remaining files
for i, filename in enumerate(FILES[1:], 2):
    src = Document(os.path.join(BASE, filename))
    
    for para in src.paragraphs:
        copy_paragraph(para, merged)
    
    print(f"  Added: {filename}")

# Add page numbers to all sections
print("Adding page numbers...")
add_page_number(merged)

# Save
merged.save(OUT)
sz = os.path.getsize(OUT)
print(f"\nSaved: {OUT}")
print(f"Size: {sz/1024:.1f} KB")
print(f"Done!")
