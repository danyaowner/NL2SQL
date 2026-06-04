import os
os.chdir("C:/Users/Danila/OneDrive/Desktop/agent")
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as A
doc = Document("coursework_text/coursework_v2_stage3.docx")
def H(t,l=0):p=doc.add_paragraph();p.alignment=A.CENTER if l==0 else A.LEFT;p.paragraph_format.space_before=Pt(18);p.paragraph_format.space_after=Pt(12);p.paragraph_format.line_spacing=1.5;r=p.add_run(t);r.font.name="Times New Roman";r.font.size=Pt(14);r.bold=True
def B(t):p=doc.add_paragraph();p.alignment=A.JUSTIFY;p.paragraph_format.first_line_indent=Cm(1.25);p.paragraph_format.space_after=Pt(6);p.paragraph_format.line_spacing=1.5;r=p.add_run(t);r.font.name="Times New Roman";r.font.size=Pt(14)
print("Functions defined, continuing...")
